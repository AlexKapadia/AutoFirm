"""Structural + determinism tests for the business-document builder.

Every generated ``.docx`` is RE-OPENED with python-docx and its real structure
asserted — the Title style, Heading-1 styled sections (a navigable outline), the
executive-summary block, every body paragraph in order, and the page-number
PAGE field in the footer — not merely that ``build`` ran. Determinism and
fail-closed paths are exercised (CLAUDE.md §3.6, §5.6).
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from hypothesis import HealthCheck, given, settings
from hypothesis import strategies as st

from autofirm.artifacts.artifact_spec_validation_errors import ArtifactSpecError
from autofirm.artifacts.business_document_builder import build_business_document
from autofirm.artifacts.business_document_spec import DocumentSection, DocumentSpec


def _doc() -> DocumentSpec:
    return DocumentSpec(
        title="Acme Annual Review",
        executive_summary="Acme grew revenue 32% and reached cashflow breakeven.",
        sections=(
            DocumentSection("Performance", ("Revenue rose to 12.4M.", "Gross margin 64%.")),
            DocumentSection("Outlook", ("FY25 guidance is 18M.",)),
        ),
    )


def _styled(path: Path) -> list[tuple[str, str]]:
    """Return ``(style_name, text)`` for each non-empty body paragraph."""
    document = Document(str(path))
    return [(p.style.name, p.text) for p in document.paragraphs if p.text.strip()]


def test_title_uses_title_style(tmp_path: Path) -> None:
    path = build_business_document(_doc(), tmp_path / "r.docx")
    styled = _styled(path)
    assert ("Title", "Acme Annual Review") in styled


def test_sections_use_heading_one_style(tmp_path: Path) -> None:
    path = build_business_document(_doc(), tmp_path / "r.docx")
    styled = _styled(path)
    headings = [text for style, text in styled if style == "Heading 1"]
    # Executive summary + the two section headings — a real outline tree.
    assert headings == ["Executive summary", "Performance", "Outlook"]


def test_body_paragraphs_present_in_order(tmp_path: Path) -> None:
    path = build_business_document(_doc(), tmp_path / "r.docx")
    body = [text for style, text in _styled(path) if style == "Normal"]
    assert body == [
        "Acme grew revenue 32% and reached cashflow breakeven.",
        "Revenue rose to 12.4M.",
        "Gross margin 64%.",
        "FY25 guidance is 18M.",
    ]


def test_executive_summary_omitted_when_blank(tmp_path: Path) -> None:
    spec = DocumentSpec(title="No Summary", sections=(DocumentSection("S", ("p",)),))
    path = build_business_document(spec, tmp_path / "r.docx")
    headings = [text for style, text in _styled(path) if style == "Heading 1"]
    assert "Executive summary" not in headings


def test_footer_contains_page_field(tmp_path: Path) -> None:
    path = build_business_document(_doc(), tmp_path / "r.docx")
    document = Document(str(path))
    footer = document.sections[0].footer.paragraphs[0]
    # The PAGE field code must be present so Word numbers pages on open.
    assert "PAGE" in footer._p.xml
    assert "Page " in footer.text


def test_document_build_is_deterministic(tmp_path: Path) -> None:
    a = build_business_document(_doc(), tmp_path / "a.docx")
    b = build_business_document(_doc(), tmp_path / "b.docx")
    assert a.read_bytes() == b.read_bytes()


@pytest.mark.parametrize("name", ["r.txt", "r.doc", "r", "report.pdf"])
def test_non_docx_destination_refused(tmp_path: Path, name: str) -> None:
    with pytest.raises(ArtifactSpecError, match=r"must be a \.docx"):
        build_business_document(_doc(), tmp_path / name)


def test_creates_missing_parent_directory(tmp_path: Path) -> None:
    path = build_business_document(_doc(), tmp_path / "nested" / "deep" / "r.docx")
    assert path.exists()


@pytest.mark.property
@settings(max_examples=30, suppress_health_check=[HealthCheck.function_scoped_fixture])
@given(
    n_sections=st.integers(min_value=1, max_value=5),
    n_paras=st.integers(min_value=1, max_value=3),
)
def test_every_section_and_paragraph_is_written(
    tmp_path: Path, n_sections: int, n_paras: int
) -> None:
    sections = tuple(
        DocumentSection(f"Heading {s}", tuple(f"Body {s}.{p}" for p in range(n_paras)))
        for s in range(n_sections)
    )
    spec = DocumentSpec(title="Generated", sections=sections)
    path = build_business_document(spec, tmp_path / "g.docx")
    styled = _styled(path)
    headings = [t for style, t in styled if style == "Heading 1"]
    body = [t for style, t in styled if style == "Normal"]
    # No section heading or body paragraph may be dropped, for any document shape.
    assert headings == [f"Heading {s}" for s in range(n_sections)]
    assert len(body) == n_sections * n_paras
