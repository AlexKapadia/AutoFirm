"""Render a structured :class:`DocumentSpec` to a clean ``.docx`` report.

What this does
--------------
Builds a real Word document from a validated document spec: a title, an optional
answer-first executive summary, then each section as a Heading-1 plus body
paragraphs, with a page-numbered footer (Minto storyline,
``docs/research/B15-artifact-generation`` §2.3). The router research recommends
python-docx for *editable* deliverables (memos/reports the client edits), which
is this builder's lane.

Why it exists / where it sits
-----------------------------
This is the L1.B15.3 document builder. It uses python-docx's built-in heading
styles (so the document carries a real outline / navigation tree) and adds a
field-coded page number in the footer that Word computes on open. Output is
deterministic for a given spec (CLAUDE.md §3.6).

Security / compliance invariants upheld
---------------------------------------
The builder refuses a non-``.docx`` target *before* writing (fail-closed —
CLAUDE.md §5.6). All content is written through python-docx text APIs, so spec
strings become document text and can never alter document structure.
"""

from __future__ import annotations

from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from autofirm.artifacts.artifact_spec_validation_errors import ArtifactSpecError

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument
    from docx.text.paragraph import Paragraph

    from autofirm.artifacts.business_document_spec import DocumentSpec

__all__ = ["build_business_document"]

_TITLE_SIZE = Pt(24)
_SUMMARY_LABEL = "Executive summary"


def build_business_document(spec: DocumentSpec, destination: Path) -> Path:
    """Write ``spec`` to ``destination`` as a clean, navigable ``.docx`` report.

    Args:
        spec: A validated document spec.
        destination: Target path; must end ``.docx``. Parent dirs are created.

    Returns:
        The ``destination`` path (now a written document).

    Raises:
        ArtifactSpecError: If ``destination`` is not a ``.docx`` path (fail-closed
            — CLAUDE.md §5.6).
    """
    if destination.suffix.lower() != ".docx":  # fail-closed: wrong container refused
        raise ArtifactSpecError(f"destination must be a .docx file, got {destination.name!r}")

    document = Document()
    _write_title(document, spec.title)
    if spec.executive_summary.strip():
        _write_executive_summary(document, spec.executive_summary)
    for section in spec.sections:
        document.add_heading(section.heading, level=1)
        for paragraph in section.paragraphs:
            document.add_paragraph(paragraph)
    _add_page_number_footer(document)

    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(destination))  # python-docx save() takes a path string
    return destination


def _write_title(document: DocxDocument, title: str) -> None:
    """Write the report title using the built-in Title style, enlarged."""
    paragraph = document.add_paragraph(title, style="Title")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if paragraph.runs:
        paragraph.runs[0].font.size = _TITLE_SIZE  # deliberate, larger than the style default


def _write_executive_summary(document: DocxDocument, summary: str) -> None:
    """Write an answer-first executive summary block (Minto)."""
    document.add_heading(_SUMMARY_LABEL, level=1)
    para = document.add_paragraph(summary)
    if para.runs:
        para.runs[0].font.italic = True  # set the summary apart from body sections


def _add_page_number_footer(document: DocxDocument) -> None:
    """Add a centred ``Page N`` footer using a Word PAGE field.

    python-docx cannot compute page counts, so we inject a field code that Word
    evaluates on open — the footer shows the real page number in every viewer.
    """
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("Page ")
    _append_page_field(paragraph)


def _append_page_field(paragraph: Paragraph) -> None:
    """Append the OOXML run sequence for a ``{ PAGE }`` field to ``paragraph``.

    Uses the raw-OOXML escape hatch (CLAUDE.md research §2.4) because python-docx
    has no high-level page-field API.
    """
    run = paragraph.add_run()
    # A Word field is fldChar(begin) -> instrText -> fldChar(end); Word replaces
    # the run text with the computed page number when the document is opened.
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)
